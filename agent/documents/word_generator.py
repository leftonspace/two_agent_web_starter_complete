"""
PHASE 2.4: Word Document Generator

Generate .docx Word documents from templates or structured data.

Features:
- Template-based generation with variable substitution (Jinja2-style)
- Structure-based generation from dictionaries
- Tables, headings, paragraphs
- Document metadata (author, title, subject)
- Formatting support

Usage:
    generator = WordDocumentGenerator()

    # From template
    generator.generate_from_template(
        template_path="templates/offer_letter.docx",
        variables={"candidate_name": "John Doe"},
        output_path="output/offer_letter.docx"
    )

    # From structure
    generator.create_from_structure(
        structure={"title": "Report", "sections": [...]},
        output_path="output/report.docx"
    )
"""

from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)


class WordDocumentGenerator:
    """
    Generate Word documents (.docx) from templates or structured data.

    Supports:
    - Template-based generation with variable substitution
    - Structure-based generation from dictionaries
    - Tables with headers
    - Headings and paragraphs
    - Document metadata
    """

    def __init__(self):
        """Initialize Word document generator"""
        self.doc = None

    def generate_from_template(
        self,
        template_path: Path,
        variables: Dict[str, Any],
        output_path: Path
    ) -> Path:
        """
        Generate .docx from template with variable substitution.

        Template uses {{variable_name}} syntax (Jinja2-style).

        Args:
            template_path: Path to template .docx file
            variables: Dictionary of variables for substitution
            output_path: Path to save generated document

        Returns:
            Path to generated document

        Raises:
            ImportError: If docxtpl not installed
            FileNotFoundError: If template not found
            Exception: If generation fails

        Example:
            generator.generate_from_template(
                Path("templates/offer_letter.docx"),
                {"candidate_name": "John Doe", "salary": "$120,000"},
                Path("output/offer.docx")
            )
        """
        try:
            from docxtpl import DocxTemplate

            template_path = Path(template_path)
            output_path = Path(output_path)

            if not template_path.exists():
                raise FileNotFoundError(f"Template not found: {template_path}")

            # Load template
            doc = DocxTemplate(str(template_path))

            # Render with variables
            doc.render(variables)

            # Ensure output directory exists
            output_path.parent.mkdir(parents=True, exist_ok=True)

            # Save
            doc.save(str(output_path))

            logger.info(f"Generated Word document from template: {output_path}")
            return output_path

        except ImportError:
            raise ImportError(
                "docxtpl is required for template-based Word generation. "
                "Install with: pip install python-docxtpl"
            )
        except Exception as e:
            logger.error(f"Failed to generate Word document from template: {e}")
            raise

    def create_from_structure(
        self,
        structure: Dict[str, Any],
        output_path: Path,
        metadata: Optional[Dict[str, str]] = None
    ) -> Path:
        """
        Create .docx from structured data.

        Structure format:
        {
            "title": "Document Title",
            "sections": [
                {
                    "heading": "Section 1",
                    "level": 1,  # 1-9
                    "content": "Paragraph text...",
                    "subsections": [...]
                }
            ],
            "tables": [
                {
                    "headers": ["Col 1", "Col 2"],
                    "rows": [["A", "B"], ["C", "D"]],
                    "style": "Light Grid Accent 1"  # Optional
                }
            ]
        }

        Args:
            structure: Document structure dictionary
            output_path: Path to save document
            metadata: Optional metadata (author, title, subject)

        Returns:
            Path to generated document

        Raises:
            ImportError: If python-docx not installed
            Exception: If generation fails

        Example:
            generator.create_from_structure(
                {
                    "title": "Quarterly Report",
                    "sections": [
                        {"heading": "Summary", "level": 1, "content": "..."}
                    ]
                },
                Path("output/report.docx")
            )
        """
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            output_path = Path(output_path)

            # Create document
            self.doc = Document()

            # Add title
            if structure.get("title"):
                title = self.doc.add_heading(structure["title"], level=0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add sections recursively
            for section in structure.get("sections", []):
                self._add_section(section)

            # Add tables
            for table_data in structure.get("tables", []):
                self._add_table(table_data)

            # Add metadata
            if metadata:
                self.add_metadata(
                    author=metadata.get("author", ""),
                    title=metadata.get("title", ""),
                    subject=metadata.get("subject", "")
                )

            # Ensure output directory exists
            output_path.parent.mkdir(parents=True, exist_ok=True)

            # Save
            self.doc.save(str(output_path))

            logger.info(f"Generated Word document from structure: {output_path}")
            return output_path

        except ImportError:
            raise ImportError(
                "python-docx is required for Word document generation. "
                "Install with: pip install python-docx"
            )
        except Exception as e:
            logger.error(f"Failed to generate Word document from structure: {e}")
            raise

    def _add_section(self, section: Dict[str, Any]) -> None:
        """
        Add section with heading and content.

        Args:
            section: Section dictionary with heading, level, content, subsections
        """
        if not self.doc:
            raise RuntimeError("Document not initialized")

        # Add heading
        if section.get("heading"):
            level = section.get("level", 1)
            self.doc.add_heading(section["heading"], level=level)

        # Add content paragraphs
        if section.get("content"):
            content = section["content"]
            if isinstance(content, str):
                self.doc.add_paragraph(content)
            elif isinstance(content, list):
                for paragraph in content:
                    self.doc.add_paragraph(str(paragraph))

        # Handle subsections recursively
        for subsection in section.get("subsections", []):
            self._add_section(subsection)

    def _add_table(self, table_data: Dict[str, Any]) -> None:
        """
        Add table with headers and rows.

        Args:
            table_data: Table dictionary with headers, rows, optional style
        """
        if not self.doc:
            raise RuntimeError("Document not initialized")

        headers = table_data.get("headers", [])
        rows = table_data.get("rows", [])
        style = table_data.get("style", "Light Grid Accent 1")

        if not headers and not rows:
            logger.warning("Skipping empty table")
            return

        # Create table
        num_rows = len(rows) + (1 if headers else 0)
        num_cols = len(headers) if headers else (len(rows[0]) if rows else 0)

        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = style

        # Add headers
        if headers:
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = str(header)

        # Add rows
        start_row = 1 if headers else 0
        for i, row in enumerate(rows):
            for j, cell_value in enumerate(row):
                table.rows[i + start_row].cells[j].text = str(cell_value)

    def add_metadata(self, author: str, title: str, subject: str) -> None:
        """
        Add document metadata.

        Args:
            author: Document author
            title: Document title
            subject: Document subject

        Example:
            generator.add_metadata(
                author="Jane Doe",
                title="Quarterly Report",
                subject="Q4 2025 Results"
            )
        """
        if not self.doc:
            raise RuntimeError("Document not initialized")

        core_props = self.doc.core_properties
        core_props.author = author
        core_props.title = title
        core_props.subject = subject
        core_props.created = datetime.now()

        logger.debug(f"Added metadata: author={author}, title={title}")


# Convenience function
def generate_word_document(
    structure: Dict[str, Any],
    output_path: Path,
    metadata: Optional[Dict[str, str]] = None
) -> Path:
    """
    Convenience function to generate Word document.

    Args:
        structure: Document structure
        output_path: Output path
        metadata: Optional metadata

    Returns:
        Path to generated document

    Example:
        generate_word_document(
            {"title": "Report", "sections": [...]},
            Path("output/report.docx")
        )
    """
    generator = WordDocumentGenerator()
    return generator.create_from_structure(structure, output_path, metadata)
