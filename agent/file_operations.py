"""
File Operations Module for JARVIS

Provides comprehensive file system access and document generation capabilities:
- Cross-platform path handling (Windows, Linux, macOS)
- PDF, Word, Excel, and text file generation
- Directory listing and file reading/writing
- External file system access beyond the project directory

This module enables JARVIS to create physical files on the user's system.
"""

from __future__ import annotations

import os
import platform
import re
from datetime import datetime
from pathlib import Path, PureWindowsPath, PurePosixPath
from typing import Any, Dict, List, Optional, Tuple, Union

# Document generation imports
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False


# =============================================================================
# Path Utilities - Cross-platform support
# =============================================================================

def normalize_path(path_str: str) -> Path:
    """
    Normalize a path string to work across platforms.

    Handles:
    - Windows paths (D:\\projects\\..., D:/projects/...)
    - Linux/Mac paths (/home/user/..., ~/projects/...)
    - Relative paths

    Args:
        path_str: Path string in any format

    Returns:
        Normalized Path object
    """
    if not path_str:
        return Path.cwd()

    # Clean up the path string
    path_str = path_str.strip().strip('"').strip("'")

    # Handle Windows-style paths on non-Windows systems
    # Check for Windows drive letter pattern (e.g., D:\, C:/, E:)
    windows_drive_pattern = r'^[A-Za-z]:[/\\]'
    if re.match(windows_drive_pattern, path_str):
        # On Windows, use as-is; on other systems, warn but try to handle
        if platform.system() == 'Windows':
            return Path(path_str)
        else:
            # Running on Linux/Mac but got Windows path
            # Try to use it as-is (might be a mounted drive or WSL path)
            return Path(path_str)

    # Handle home directory expansion
    if path_str.startswith('~'):
        return Path(path_str).expanduser()

    # Handle relative paths
    if not os.path.isabs(path_str):
        return Path.cwd() / path_str

    return Path(path_str)


def is_valid_path(path: Union[str, Path]) -> Tuple[bool, str]:
    """
    Check if a path is valid and accessible.

    Args:
        path: Path to validate

    Returns:
        Tuple of (is_valid, message)
    """
    try:
        if isinstance(path, str):
            path = normalize_path(path)

        # Check if parent directory exists (for new files)
        parent = path.parent
        if parent.exists():
            return True, "Path is valid"
        else:
            return False, f"Parent directory does not exist: {parent}"

    except Exception as e:
        return False, f"Invalid path: {e}"


def ensure_directory(path: Union[str, Path]) -> Tuple[bool, str]:
    """
    Ensure a directory exists, creating it if necessary.

    Args:
        path: Directory path to ensure exists

    Returns:
        Tuple of (success, message)
    """
    try:
        if isinstance(path, str):
            path = normalize_path(path)

        path.mkdir(parents=True, exist_ok=True)
        return True, f"Directory ready: {path}"

    except PermissionError:
        return False, f"Permission denied creating directory: {path}"
    except Exception as e:
        return False, f"Error creating directory: {e}"


# =============================================================================
# File System Operations
# =============================================================================

def list_directory(
    path: Union[str, Path],
    recursive: bool = False,
    include_hidden: bool = False,
    file_types: Optional[List[str]] = None
) -> Dict[str, Any]:
    """
    List contents of a directory.

    Args:
        path: Directory path to list
        recursive: Whether to list subdirectories recursively
        include_hidden: Include hidden files/directories
        file_types: Filter by file extensions (e.g., ['.py', '.txt'])

    Returns:
        Dict with directory contents and metadata
    """
    try:
        if isinstance(path, str):
            path = normalize_path(path)

        if not path.exists():
            return {
                "success": False,
                "error": f"Path does not exist: {path}",
                "path": str(path),
                "files": [],
                "directories": []
            }

        if not path.is_dir():
            return {
                "success": False,
                "error": f"Path is not a directory: {path}",
                "path": str(path),
                "files": [],
                "directories": []
            }

        files = []
        directories = []

        def should_include(item: Path) -> bool:
            if not include_hidden and item.name.startswith('.'):
                return False
            if file_types and item.is_file():
                return item.suffix.lower() in [ext.lower() for ext in file_types]
            return True

        if recursive:
            for item in path.rglob('*'):
                if should_include(item):
                    if item.is_file():
                        files.append({
                            "name": item.name,
                            "path": str(item),
                            "relative_path": str(item.relative_to(path)),
                            "size": item.stat().st_size,
                            "modified": datetime.fromtimestamp(item.stat().st_mtime).isoformat()
                        })
                    elif item.is_dir():
                        directories.append({
                            "name": item.name,
                            "path": str(item),
                            "relative_path": str(item.relative_to(path))
                        })
        else:
            for item in path.iterdir():
                if should_include(item):
                    if item.is_file():
                        files.append({
                            "name": item.name,
                            "path": str(item),
                            "size": item.stat().st_size,
                            "modified": datetime.fromtimestamp(item.stat().st_mtime).isoformat()
                        })
                    elif item.is_dir():
                        directories.append({
                            "name": item.name,
                            "path": str(item)
                        })

        return {
            "success": True,
            "path": str(path),
            "files": sorted(files, key=lambda x: x["name"]),
            "directories": sorted(directories, key=lambda x: x["name"]),
            "total_files": len(files),
            "total_directories": len(directories)
        }

    except PermissionError:
        return {
            "success": False,
            "error": f"Permission denied accessing: {path}",
            "path": str(path),
            "files": [],
            "directories": []
        }
    except Exception as e:
        return {
            "success": False,
            "error": f"Error listing directory: {e}",
            "path": str(path) if isinstance(path, Path) else path,
            "files": [],
            "directories": []
        }


def read_file(
    path: Union[str, Path],
    encoding: str = 'utf-8'
) -> Dict[str, Any]:
    """
    Read content from a file.

    Args:
        path: File path to read
        encoding: Text encoding (default: utf-8)

    Returns:
        Dict with file content and metadata
    """
    try:
        if isinstance(path, str):
            path = normalize_path(path)

        if not path.exists():
            return {
                "success": False,
                "error": f"File does not exist: {path}",
                "path": str(path),
                "content": None
            }

        if not path.is_file():
            return {
                "success": False,
                "error": f"Path is not a file: {path}",
                "path": str(path),
                "content": None
            }

        # Check file size
        size = path.stat().st_size
        if size > 10 * 1024 * 1024:  # 10MB limit
            return {
                "success": False,
                "error": f"File too large ({size / 1024 / 1024:.1f} MB). Maximum is 10 MB.",
                "path": str(path),
                "content": None,
                "size": size
            }

        # Try to read as text
        try:
            content = path.read_text(encoding=encoding)
            return {
                "success": True,
                "path": str(path),
                "content": content,
                "size": size,
                "encoding": encoding,
                "is_binary": False
            }
        except UnicodeDecodeError:
            # File is binary
            return {
                "success": True,
                "path": str(path),
                "content": None,
                "size": size,
                "is_binary": True,
                "error": "File is binary and cannot be displayed as text"
            }

    except PermissionError:
        return {
            "success": False,
            "error": f"Permission denied reading: {path}",
            "path": str(path),
            "content": None
        }
    except Exception as e:
        return {
            "success": False,
            "error": f"Error reading file: {e}",
            "path": str(path) if isinstance(path, Path) else path,
            "content": None
        }


def write_file(
    path: Union[str, Path],
    content: str,
    encoding: str = 'utf-8',
    create_dirs: bool = True
) -> Dict[str, Any]:
    """
    Write content to a text file.

    Args:
        path: File path to write
        content: Content to write
        encoding: Text encoding (default: utf-8)
        create_dirs: Create parent directories if they don't exist

    Returns:
        Dict with success status and details
    """
    try:
        if isinstance(path, str):
            path = normalize_path(path)

        # Create parent directories if needed
        if create_dirs:
            path.parent.mkdir(parents=True, exist_ok=True)

        path.write_text(content, encoding=encoding)

        return {
            "success": True,
            "path": str(path),
            "size": path.stat().st_size,
            "message": f"File created successfully: {path}"
        }

    except PermissionError:
        return {
            "success": False,
            "error": f"Permission denied writing to: {path}",
            "path": str(path)
        }
    except Exception as e:
        return {
            "success": False,
            "error": f"Error writing file: {e}",
            "path": str(path) if isinstance(path, Path) else path
        }


# =============================================================================
# Document Generation
# =============================================================================

def create_pdf(
    path: Union[str, Path],
    content: str,
    title: Optional[str] = None,
    author: str = "JARVIS",
    page_size: str = "letter"
) -> Dict[str, Any]:
    """
    Create a PDF document.

    Args:
        path: Output file path
        content: Text content for the PDF
        title: Document title (shown at top)
        author: Document author
        page_size: 'letter' or 'A4'

    Returns:
        Dict with success status and file details
    """
    if not REPORTLAB_AVAILABLE:
        return {
            "success": False,
            "error": "PDF generation not available. Please install reportlab: pip install reportlab",
            "path": str(path) if isinstance(path, Path) else path
        }

    try:
        if isinstance(path, str):
            path = normalize_path(path)

        # Ensure .pdf extension
        if path.suffix.lower() != '.pdf':
            path = path.with_suffix('.pdf')

        # Create parent directories
        path.parent.mkdir(parents=True, exist_ok=True)

        # Set page size
        size = letter if page_size.lower() == 'letter' else A4

        # Create PDF
        doc = SimpleDocTemplate(
            str(path),
            pagesize=size,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )

        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            alignment=TA_CENTER,
            spaceAfter=20
        )

        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            alignment=TA_JUSTIFY,
            spaceBefore=6,
            spaceAfter=6,
            leading=14
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceBefore=12,
            spaceAfter=8
        )

        # Build document content
        story = []

        # Add title if provided
        if title:
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 12))

        # Add timestamp
        timestamp = datetime.now().strftime("%B %d, %Y at %H:%M")
        story.append(Paragraph(f"Generated by {author} on {timestamp}", styles['Italic']))
        story.append(Spacer(1, 20))

        # Process content - handle markdown-like formatting
        lines = content.split('\n')
        for line in lines:
            line = line.strip()

            if not line:
                story.append(Spacer(1, 6))
            elif line.startswith('# '):
                story.append(Paragraph(line[2:], title_style))
            elif line.startswith('## '):
                story.append(Paragraph(line[3:], heading_style))
            elif line.startswith('### '):
                story.append(Paragraph(line[4:], heading_style))
            elif line.startswith('- ') or line.startswith('* '):
                story.append(Paragraph(f"• {line[2:]}", body_style))
            else:
                # Escape special characters for reportlab
                safe_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(safe_line, body_style))

        # Build PDF
        doc.build(story)

        return {
            "success": True,
            "path": str(path),
            "size": path.stat().st_size,
            "message": f"PDF created successfully: {path}"
        }

    except Exception as e:
        return {
            "success": False,
            "error": f"Error creating PDF: {e}",
            "path": str(path) if isinstance(path, Path) else path
        }


def create_word_document(
    path: Union[str, Path],
    content: str,
    title: Optional[str] = None,
    author: str = "JARVIS"
) -> Dict[str, Any]:
    """
    Create a Word document (.docx).

    Args:
        path: Output file path
        content: Text content for the document
        title: Document title
        author: Document author

    Returns:
        Dict with success status and file details
    """
    if not DOCX_AVAILABLE:
        return {
            "success": False,
            "error": "Word document generation not available. Please install python-docx: pip install python-docx",
            "path": str(path) if isinstance(path, Path) else path
        }

    try:
        if isinstance(path, str):
            path = normalize_path(path)

        # Ensure .docx extension
        if path.suffix.lower() != '.docx':
            path = path.with_suffix('.docx')

        # Create parent directories
        path.parent.mkdir(parents=True, exist_ok=True)

        # Create document
        doc = Document()

        # Set document properties
        doc.core_properties.author = author
        doc.core_properties.created = datetime.now()

        # Add title if provided
        if title:
            title_para = doc.add_heading(title, 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add timestamp
        timestamp = datetime.now().strftime("%B %d, %Y at %H:%M")
        meta_para = doc.add_paragraph(f"Generated by {author} on {timestamp}")
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_para.runs[0].italic = True

        doc.add_paragraph()  # Spacer

        # Process content - handle markdown-like formatting
        lines = content.split('\n')
        for line in lines:
            line = line.strip()

            if not line:
                doc.add_paragraph()
            elif line.startswith('# '):
                doc.add_heading(line[2:], 1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], 2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], 3)
            elif line.startswith('- ') or line.startswith('* '):
                para = doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
                para = doc.add_paragraph(line[3:], style='List Number')
            else:
                doc.add_paragraph(line)

        # Save document
        doc.save(str(path))

        return {
            "success": True,
            "path": str(path),
            "size": path.stat().st_size,
            "message": f"Word document created successfully: {path}"
        }

    except Exception as e:
        return {
            "success": False,
            "error": f"Error creating Word document: {e}",
            "path": str(path) if isinstance(path, Path) else path
        }


def create_excel_spreadsheet(
    path: Union[str, Path],
    data: List[List[Any]],
    headers: Optional[List[str]] = None,
    sheet_name: str = "Sheet1",
    title: Optional[str] = None
) -> Dict[str, Any]:
    """
    Create an Excel spreadsheet (.xlsx).

    Args:
        path: Output file path
        data: 2D list of data rows
        headers: Column headers
        sheet_name: Name of the worksheet
        title: Document title (added above data)

    Returns:
        Dict with success status and file details
    """
    if not OPENPYXL_AVAILABLE:
        return {
            "success": False,
            "error": "Excel generation not available. Please install openpyxl: pip install openpyxl",
            "path": str(path) if isinstance(path, Path) else path
        }

    try:
        if isinstance(path, str):
            path = normalize_path(path)

        # Ensure .xlsx extension
        if path.suffix.lower() != '.xlsx':
            path = path.with_suffix('.xlsx')

        # Create parent directories
        path.parent.mkdir(parents=True, exist_ok=True)

        # Create workbook
        wb = Workbook()
        ws = wb.active
        ws.title = sheet_name

        start_row = 1

        # Add title if provided
        if title:
            ws.cell(row=1, column=1, value=title)
            ws.cell(row=1, column=1).font = Font(bold=True, size=14)
            ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers) if headers else len(data[0]) if data else 1)
            start_row = 3

        # Add headers if provided
        if headers:
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=start_row, column=col, value=header)
                cell.font = Font(bold=True)
                cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
                cell.alignment = Alignment(horizontal="center")
            start_row += 1

        # Add data
        for row_idx, row_data in enumerate(data, start_row):
            for col_idx, value in enumerate(row_data, 1):
                ws.cell(row=row_idx, column=col_idx, value=value)

        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if cell.value:
                        max_length = max(max_length, len(str(cell.value)))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width

        # Save workbook
        wb.save(str(path))

        return {
            "success": True,
            "path": str(path),
            "size": path.stat().st_size,
            "rows": len(data),
            "columns": len(headers) if headers else (len(data[0]) if data else 0),
            "message": f"Excel spreadsheet created successfully: {path}"
        }

    except Exception as e:
        return {
            "success": False,
            "error": f"Error creating Excel spreadsheet: {e}",
            "path": str(path) if isinstance(path, Path) else path
        }


def create_markdown_file(
    path: Union[str, Path],
    content: str,
    title: Optional[str] = None
) -> Dict[str, Any]:
    """
    Create a Markdown file.

    Args:
        path: Output file path
        content: Markdown content
        title: Optional title to prepend

    Returns:
        Dict with success status and file details
    """
    try:
        if isinstance(path, str):
            path = normalize_path(path)

        # Ensure .md extension
        if path.suffix.lower() not in ['.md', '.markdown']:
            path = path.with_suffix('.md')

        # Create parent directories
        path.parent.mkdir(parents=True, exist_ok=True)

        # Build content
        full_content = ""
        if title:
            full_content = f"# {title}\n\n"
        full_content += content

        # Add generation timestamp at bottom
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
        full_content += f"\n\n---\n*Generated by JARVIS on {timestamp}*\n"

        # Write file
        path.write_text(full_content, encoding='utf-8')

        return {
            "success": True,
            "path": str(path),
            "size": path.stat().st_size,
            "message": f"Markdown file created successfully: {path}"
        }

    except Exception as e:
        return {
            "success": False,
            "error": f"Error creating Markdown file: {e}",
            "path": str(path) if isinstance(path, Path) else path
        }


# =============================================================================
# Convenience Functions
# =============================================================================

def create_document(
    path: Union[str, Path],
    content: str,
    format: str = "auto",
    title: Optional[str] = None,
    **kwargs
) -> Dict[str, Any]:
    """
    Create a document in the specified format.

    Args:
        path: Output file path
        content: Document content
        format: 'pdf', 'docx', 'xlsx', 'md', 'txt', or 'auto' (detect from extension)
        title: Document title
        **kwargs: Additional format-specific arguments

    Returns:
        Dict with success status and file details
    """
    if isinstance(path, str):
        path = normalize_path(path)

    # Auto-detect format from extension
    if format == "auto":
        ext = path.suffix.lower()
        format_map = {
            '.pdf': 'pdf',
            '.docx': 'docx',
            '.doc': 'docx',
            '.xlsx': 'xlsx',
            '.xls': 'xlsx',
            '.md': 'md',
            '.markdown': 'md',
            '.txt': 'txt'
        }
        format = format_map.get(ext, 'txt')

    # Route to appropriate generator
    if format == 'pdf':
        return create_pdf(path, content, title=title, **kwargs)
    elif format == 'docx':
        return create_word_document(path, content, title=title, **kwargs)
    elif format == 'xlsx':
        # For Excel, content should be parsed or passed as data
        if 'data' in kwargs:
            return create_excel_spreadsheet(path, kwargs['data'], title=title, **kwargs)
        else:
            # Convert text content to simple rows
            lines = [line.strip() for line in content.split('\n') if line.strip()]
            data = [[line] for line in lines]
            return create_excel_spreadsheet(path, data, title=title, **kwargs)
    elif format == 'md':
        return create_markdown_file(path, content, title=title)
    else:  # txt
        return write_file(path, content if not title else f"{title}\n\n{content}")


def get_available_formats() -> Dict[str, bool]:
    """
    Get available document formats and their availability status.

    Returns:
        Dict mapping format names to availability booleans
    """
    return {
        "pdf": REPORTLAB_AVAILABLE,
        "docx": DOCX_AVAILABLE,
        "xlsx": OPENPYXL_AVAILABLE,
        "md": True,  # Always available
        "txt": True  # Always available
    }


# =============================================================================
# File Operation Handler for JARVIS
# =============================================================================

class FileOperationHandler:
    """
    High-level file operation handler for JARVIS integration.

    Provides natural language-friendly interface for file operations.
    """

    def __init__(self, default_output_dir: Optional[str] = None):
        """
        Initialize the file operation handler.

        Args:
            default_output_dir: Default directory for file outputs
        """
        self.default_output_dir = normalize_path(default_output_dir) if default_output_dir else Path.cwd()

    def get_capabilities(self) -> str:
        """Get a description of available capabilities."""
        formats = get_available_formats()
        available = [fmt for fmt, avail in formats.items() if avail]

        return f"""File Operation Capabilities:
- Create files: {', '.join(available).upper()}
- Read files from any accessible path
- List directory contents
- Create directories
- Cross-platform path support (Windows, Linux, macOS)"""

    def handle_request(
        self,
        operation: str,
        path: str,
        content: Optional[str] = None,
        **kwargs
    ) -> Dict[str, Any]:
        """
        Handle a file operation request.

        Args:
            operation: 'create', 'read', 'list', 'mkdir'
            path: Target path
            content: Content for create operations
            **kwargs: Additional arguments

        Returns:
            Operation result dictionary
        """
        operation = operation.lower().strip()

        if operation == 'create':
            if not content:
                return {"success": False, "error": "Content required for create operation"}
            return create_document(path, content, **kwargs)

        elif operation == 'read':
            return read_file(path)

        elif operation == 'list':
            return list_directory(path, **kwargs)

        elif operation == 'mkdir':
            success, message = ensure_directory(path)
            return {"success": success, "message": message, "path": str(normalize_path(path))}

        else:
            return {"success": False, "error": f"Unknown operation: {operation}"}

    def create_from_previous_output(
        self,
        content: str,
        format: str,
        filename: str,
        output_dir: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Create a file from previous JARVIS output.

        Args:
            content: The content to save
            format: Output format ('pdf', 'docx', 'md', 'txt')
            filename: Base filename (without extension)
            output_dir: Output directory (uses default if not specified)

        Returns:
            Operation result dictionary
        """
        output_path = normalize_path(output_dir) if output_dir else self.default_output_dir

        # Build full path
        ext_map = {
            'pdf': '.pdf',
            'docx': '.docx',
            'word': '.docx',
            'xlsx': '.xlsx',
            'excel': '.xlsx',
            'md': '.md',
            'markdown': '.md',
            'txt': '.txt',
            'text': '.txt'
        }

        ext = ext_map.get(format.lower(), '.txt')
        full_path = output_path / f"{filename}{ext}"

        # Map format names
        format_map = {
            'word': 'docx',
            'excel': 'xlsx',
            'markdown': 'md',
            'text': 'txt'
        }
        actual_format = format_map.get(format.lower(), format.lower())

        return create_document(full_path, content, format=actual_format)


# =============================================================================
# Exports
# =============================================================================

__all__ = [
    # Path utilities
    'normalize_path',
    'is_valid_path',
    'ensure_directory',

    # File system operations
    'list_directory',
    'read_file',
    'write_file',

    # Document generation
    'create_pdf',
    'create_word_document',
    'create_excel_spreadsheet',
    'create_markdown_file',
    'create_document',

    # Utilities
    'get_available_formats',

    # Handler class
    'FileOperationHandler',
]
